#!/usr/bin/env python3
"""Ocean Eterna Document Preprocessor — converts PDF, DOCX, XLSX, CSV, images to indexed text.

Usage:
    python3 doc_processor.py report.pdf                     # ingest into OE
    python3 doc_processor.py *.docx                         # batch ingest
    python3 doc_processor.py --dry-run scan.png             # preview extraction
    python3 doc_processor.py --server http://host:9090 f.pdf  # custom server
"""
import os
import re
import csv
import sys
import argparse
from pathlib import Path

import requests

OE_BASE = os.environ.get("OE_BASE_URL", "http://localhost:9090")
MAX_CONTENT_SIZE = 10_000_000  # 10MB text cap
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".csv", ".png", ".jpg", ".jpeg", ".txt", ".md"}


def normalize_paragraphs(text: str) -> str:
    """clean text to UTF-8 with consistent \\n\\n paragraph separators."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\S\n]+", " ", text)
    text = re.sub(r" +\n", "\n", text)
    return text.strip()


# ── extractors ───────────────────────────────────────────────────────

def extract_pdf(filepath: str) -> str:
    try:
        import pymupdf
    except ImportError:
        raise ImportError("PDF support requires pymupdf: pip install pymupdf")

    try:
        doc = pymupdf.open(filepath)
    except Exception as e:
        raise ValueError(f"cannot open PDF '{filepath}': {e}")

    pages = []
    for page_num, page in enumerate(doc):
        text = page.get_text("text").strip()

        # scanned page — try OCR fallback
        if len(text) < 50:
            try:
                import pytesseract
                from PIL import Image

                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img, lang="eng").strip()
                pix = None
            except Exception:
                if not text:
                    text = f"[Page {page_num + 1}: could not extract text]"

        if text:
            pages.append(f"[Page {page_num + 1}]\n\n{text}")

    doc.close()
    return "\n\n".join(pages)


def extract_docx(filepath: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("DOCX support requires python-docx: pip install python-docx")

    try:
        doc = Document(filepath)
    except Exception as e:
        raise ValueError(f"cannot open DOCX '{filepath}': {e}")

    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "").strip()
            prefix = "#" * int(level) if level.isdigit() else "##"
            parts.append(f"{prefix} {text}")
        else:
            parts.append(text)

    # extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def extract_xlsx(filepath: str) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("XLSX support requires openpyxl: pip install openpyxl")

    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    except Exception as e:
        raise ValueError(f"cannot open XLSX '{filepath}': {e}")

    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            header = f"## Sheet: {sheet_name}"
            # batch rows in groups of 50 for paragraph boundaries
            batched = []
            for i in range(0, len(rows), 50):
                batched.append("\n".join(rows[i : i + 50]))
            parts.append(f"{header}\n\n" + "\n\n".join(batched))

    wb.close()
    return "\n\n".join(parts)


def extract_csv(filepath: str) -> str:
    try:
        with open(filepath, newline="", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = []
            for row in reader:
                if any(cell.strip() for cell in row):
                    rows.append(" | ".join(row))
    except Exception as e:
        raise ValueError(f"cannot read CSV '{filepath}': {e}")

    if not rows:
        return ""

    # batch rows in groups of 50
    parts = []
    for i in range(0, len(rows), 50):
        parts.append("\n".join(rows[i : i + 50]))
    return "\n\n".join(parts)


def extract_image(filepath: str) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise ImportError(
            "Image OCR requires pytesseract and Pillow: "
            "pip install pytesseract pillow && sudo apt install tesseract-ocr"
        )

    try:
        pytesseract.get_tesseract_version()
    except Exception:
        raise RuntimeError(
            "Tesseract OCR binary not found. Install: sudo apt install tesseract-ocr"
        )

    img = Image.open(filepath)
    if img.mode not in ("L", "RGB"):
        img = img.convert("RGB")

    text = pytesseract.image_to_string(img, lang="eng")
    img.close()
    return text


def extract_plaintext(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


# ── dispatcher ───────────────────────────────────────────────────────

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".csv": extract_csv,
    ".png": extract_image,
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".txt": extract_plaintext,
    ".md": extract_plaintext,
}


def process_document(filepath: str) -> tuple:
    """extract text from a document file.

    returns (filename, content) tuple where content is clean UTF-8 text
    with \\n\\n paragraph separators, ready for OE ingestion.
    """
    path = Path(filepath)

    if not path.exists():
        raise FileNotFoundError(f"file not found: {filepath}")
    if not path.is_file():
        raise ValueError(f"not a file: {filepath}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"unsupported format '{ext}'. supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    raw_text = EXTRACTORS[ext](str(path))
    content = normalize_paragraphs(raw_text)

    if not content.strip():
        raise ValueError(f"no text content extracted from {filepath}")

    if len(content) > MAX_CONTENT_SIZE:
        content = content[:MAX_CONTENT_SIZE]
        content += "\n\n[Document truncated at 10MB text limit]"

    return (path.name, content)


# ── server communication ─────────────────────────────────────────────

def send_to_oe(filename: str, content: str, server_url: str = OE_BASE) -> dict:
    """send extracted text to OE server for ingestion."""
    try:
        r = requests.post(
            f"{server_url}/add-file",
            json={"filename": filename, "content": content},
            timeout=120,
        )
        return r.json()
    except requests.ConnectionError:
        return {"error": f"cannot connect to OE server at {server_url}"}
    except requests.Timeout:
        return {"error": "request timed out (file may be very large)"}
    except Exception as e:
        return {"error": str(e)}


# ── CLI ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Ocean Eterna Document Preprocessor — convert documents to indexed text"
    )
    parser.add_argument("input", nargs="+", help="file path(s) to process")
    parser.add_argument(
        "--server",
        default=os.environ.get("OE_BASE_URL", "http://localhost:9090"),
        help="OE server URL (default: http://localhost:9090, env: OE_BASE_URL)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="extract text and print to stdout, don't send to server",
    )
    parser.add_argument(
        "--output",
        help="write extracted text to file instead of sending to server",
    )
    args = parser.parse_args()

    success = 0
    errors = 0

    for filepath in args.input:
        print(f"\nProcessing: {filepath}")
        try:
            filename, content = process_document(filepath)
            chars = len(content)
            paragraphs = content.count("\n\n") + 1
            print(f"  extracted: {chars:,} chars, ~{paragraphs} paragraphs")

            if args.dry_run:
                print(f"\n--- {filename} ---")
                print(content[:2000])
                if len(content) > 2000:
                    print(f"\n... [{chars - 2000:,} more chars] ...")
                success += 1

            elif args.output:
                with open(args.output, "a", encoding="utf-8") as f:
                    f.write(f"\n\n=== {filename} ===\n\n{content}")
                print(f"  written to: {args.output}")
                success += 1

            else:
                result = send_to_oe(filename, content, args.server)
                if result.get("success"):
                    chunks = result.get("chunks_added", 0)
                    tokens = result.get("tokens_added", 0)
                    print(f"  ingested: {chunks} chunks, {tokens:,} tokens")
                    success += 1
                else:
                    print(f"  ERROR: {result.get('error', 'unknown error')}")
                    errors += 1

        except Exception as e:
            print(f"  ERROR: {e}")
            errors += 1

    print(f"\nDone: {success} succeeded, {errors} failed")
    if errors > 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
